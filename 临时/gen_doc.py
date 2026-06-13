from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 标题
title = doc.add_heading('Bigseller助手 使用说明', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 简介
doc.add_heading('简介', level=1)
doc.add_paragraph(
    'Bigseller助手是一个油猴脚本，用于在 Bigseller 订单处理页面快速执行换货操作。'
    '支持智能推荐同款SKU、一键换货、换货记录云端同步。'
)

# 安装
doc.add_heading('安装', level=1)
doc.add_paragraph('1. 安装 Tampermonkey 浏览器扩展', style='List Number')
doc.add_paragraph('2. 将 Bigseller助手.txt 中的代码复制到 Tampermonkey 新建脚本中保存', style='List Number')
doc.add_paragraph('3. 首次使用时浏览器会弹出"允许跨域请求"提示，点击允许', style='List Number')

# 适用页面
doc.add_heading('适用页面', level=1)
doc.add_paragraph('https://www.bigseller.pro/web/order/index.htm?status=new（待处理订单页）')

# 功能说明
doc.add_heading('功能说明', level=1)

doc.add_heading('1. 换货按钮', level=2)
doc.add_paragraph('页面加载后，每个订单的 SKU 编号旁会出现一个紫色的「换货」按钮。')

doc.add_heading('2. 换货弹窗', level=2)
doc.add_paragraph('点击「换货」按钮后弹出换货操作面板，包含：')
doc.add_paragraph('当前SKU（左侧）：显示商品图片、SKU编号、商品标题、可用库存', style='List Bullet')
doc.add_paragraph('目标SKU（右侧）：输入目标SKU编号，失焦后自动查询图片、标题、库存', style='List Bullet')
doc.add_paragraph('推荐SKU列表（下方）：自动搜索同款SKU并按优先级排序', style='List Bullet')

doc.add_heading('3. 推荐SKU排序规则', level=2)
doc.add_paragraph('推荐列表按以下优先级排序（有库存的始终排在前面）：')

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = table.rows[0].cells
headers[0].text = '优先级'
headers[1].text = '标签'
headers[2].text = '含义'
row1 = table.rows[1].cells
row1[0].text = '1'
row1[1].text = '异图同尺'
row1[2].text = '同尺寸、不同图案（宽高完全一致）'
row2 = table.rows[2].cells
row2[0].text = '2'
row2[1].text = '同图更大'
row2[2].text = '同图案、更大尺寸'
row3 = table.rows[3].cells
row3[0].text = '3'
row3[1].text = '异图更大'
row3[2].text = '不同图案、更大尺寸'

doc.add_paragraph('')
doc.add_paragraph('同图更大/异图更大内部：优先高度相等宽度更大的，其次面积越小越靠前。')

doc.add_heading('4. 确认换货', level=2)
doc.add_paragraph('点击推荐列表中的SKU或手动输入目标SKU后，点击「确认换货」：')
doc.add_paragraph('查询目标SKU的 skuId', style='List Number')
doc.add_paragraph('获取订单内部 orderId', style='List Number')
doc.add_paragraph('获取订单商品项 orderItemId', style='List Number')
doc.add_paragraph('调用换货接口完成SKU映射更换', style='List Number')
doc.add_paragraph('换货成功后在页面原SKU下方显示绿色结果卡片', style='List Number')
doc.add_paragraph('同时将换货记录保存到云端', style='List Number')

doc.add_heading('5. 换货记录云端同步', level=2)
doc.add_paragraph('每次换货成功后，记录自动推送到聚树ERP云端数据库', style='List Bullet')
doc.add_paragraph('页面加载时，自动从云端拉取当前页订单的换货记录', style='List Bullet')
doc.add_paragraph('已换货的订单会在SKU下方显示绿色卡片（箭头 + 目标SKU信息）', style='List Bullet')

# 操作流程
doc.add_heading('操作流程', level=1)
doc.add_paragraph('打开 Bigseller 待处理订单页', style='List Number')
doc.add_paragraph('找到需要换货的订单，点击SKU旁的紫色「换货」按钮', style='List Number')
doc.add_paragraph('在弹窗中查看推荐列表，点击合适的SKU（或手动输入）', style='List Number')
doc.add_paragraph('确认目标SKU的标题、图片、库存无误', style='List Number')
doc.add_paragraph('点击「确认换货」', style='List Number')
doc.add_paragraph('看到"换货成功"提示即完成', style='List Number')

# 注意事项
doc.add_heading('注意事项', level=1)
doc.add_paragraph('换货仅限同款（同SKU前缀），推荐列表已自动过滤', style='List Bullet')
doc.add_paragraph('换货操作不可撤销，确认前请仔细核对目标SKU', style='List Bullet')
doc.add_paragraph('仅对当前订单生效（changeType=only）', style='List Bullet')
doc.add_paragraph('弹窗打开时背景页面锁定滚动，关闭后恢复', style='List Bullet')
doc.add_paragraph('推荐列表最多显示500条结果，超长时可滚动查看', style='List Bullet')

doc.save(r'C:\Users\fjh\Desktop\油猴脚本\Bigseller相关\Bigseller助手使用说明.docx')
print('done')
